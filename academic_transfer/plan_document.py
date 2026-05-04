# -*- coding: utf-8 -*-
"""Формирование индивидуального плана (.docx) точно по шаблону ЯГТУ"""
from __future__ import annotations

import os
import re
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


# ============================================================================
# КОНСТАНТЫ
# ============================================================================

DEFAULT_PLAN_META = {
    'student_genitive': 'Егорова Никиты Сергеевича',
    'student_sign_fio': 'Егоров Н.С.',
    'direction': '09.03.02 «Информационные системы и технологии»',
    'group': 'ЦИС-36',
    'profile': 'Информационные системы и технологии',
    'deadline': 'до «1» февраля 2027 г.',
    'sign_year': '2026',
    'deputy_title': 'Заместитель директора Института цифровых систем',
    'deputy_fio': 'Бойков С.Ю.',
    'prorector_fio': 'В.А. Голкина',
}

PLAN_META_KEYS = list(DEFAULT_PLAN_META.keys())


# ============================================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ============================================================================

def _run_font(run, size_pt=12, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    run.bold = bold


def _para(doc, text, bold=False, align=None, size_pt=12):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    _run_font(r, size_pt=size_pt, bold=bold)
    return p


def _cell_write(cell, text, size_pt=10, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(str(text) if text is not None else '')
    _run_font(r, size_pt=size_pt, bold=bold)


def _fmt_ze(h):
    if h is None or h == '':
        return '—'
    try:
        if float(h) == int(float(h)):
            return str(int(float(h)))
        return str(h)
    except (TypeError, ValueError):
        return str(h)


def _fmt_sem(v) -> str:
    if v is None or v == '':
        return '—'
    try:
        if isinstance(v, float) and not v == int(v):
            return str(v)
        if float(v) == int(float(v)):
            return str(int(float(v)))
        return str(v)
    except (TypeError, ValueError):
        s = str(v).strip()
        return s if s else '—'


def _fmt_grade_cell(grade: Any, normalized: Any = None) -> str:
    if normalized is not None:
        if normalized in (5, 4, 3, 2):
            if normalized == 5:
                return 'отлично'
            elif normalized == 4:
                return 'хорошо'
            elif normalized == 3:
                return 'удовлетворительно'
            elif normalized == 2:
                return 'неудовлетворительно'
        if normalized == 'зачет':
            return 'зачтено'
        if normalized == 'не зачет':
            return 'не зачтено'
    if grade is None or grade == '':
        return '—'
    s = str(grade).strip()
    sl = s.lower()
    if sl in ('не сдано', 'не зачтено', 'не зачет') or 'не зач' in sl:
        return 'не зачтено'
    if 'удовлетвор' in sl:
        return 'удовлетворительно'
    if 'хорошо' in sl:
        return 'хорошо'
    if 'отличн' in sl:
        return 'отлично'
    if 'зачт' in sl or sl == 'зачет':
        return 'зачтено'
    if s.isdigit():
        if s == '5':
            return 'отлично'
        elif s == '4':
            return 'хорошо'
        elif s == '3':
            return 'удовлетворительно'
        elif s == '2':
            return 'неудовлетворительно'
        return s
    return s


def _get_control_form(discipline: Dict) -> str:
    """Определение формы контроля по дисциплине"""
    grade = discipline.get('grade', '')
    normalized = discipline.get('normalized_grade', '')
    control = discipline.get('control_form', '')
    
    if control:
        return control
    
    grade_str = str(grade).lower()
    if 'экзамен' in grade_str:
        return 'экзамен'
    elif 'диф.зачет' in grade_str or 'дифференцированный' in grade_str:
        return 'диф.зачет'
    elif 'курсовая' in grade_str or 'курсовой проект' in grade_str:
        return 'курсовая работа'
    elif 'зачет' in grade_str:
        return 'зачет'
    elif normalized in (5, 4, 3):
        return 'экзамен'
    elif normalized == 'зачет':
        return 'зачет'
    
    return '—'


def _get_cell_by_headers(row, headers_map):
    """
    Получение ячейки по названию заголовка (для шаблона с нестандартной структурой)
    """
    result = {}
    for idx, cell in enumerate(row.cells):
        cell_text = cell.text.strip().lower()
        for header_key, header_name in headers_map.items():
            if header_key in cell_text or cell_text == header_name.lower():
                result[header_name] = idx
                break
    return result


# ============================================================================
# ЗАПОЛНЕНИЕ ТАБЛИЦ ПО ШАБЛОНУ
# ============================================================================

def fill_table_1_old_plan(table, transcript_disciplines: List[Dict], meta: Dict):
    """
    Таблица 1: Дисциплины учебного плана ЯГТУ
    Структура: 
    - Первые строки: заголовки с информацией о направлении и институте
    - Строка с заголовками колонок: Название | Семестр | Объем, ЗЕ | Форма контроля
    - Последний столбец: Оценка (отдельный заголовок справа)
    
    В шаблоне таблица может иметь объединённые ячейки в первых строках
    """
    if table is None:
        return
    
    # Определяем, сколько строк в таблице - заголовочных
    # В шаблоне первые 2-3 строки - это информация о направлении и институте
    
    # Сохраняем заголовочные строки (первые 2-3 строки)
    header_rows = []
    data_start_row = 0
    
    for i, row in enumerate(table.rows):
        row_text = ' '.join(cell.text for cell in row.cells).lower()
        # Ищем строку с заголовками колонок
        if 'название дисциплины' in row_text or 'наименование' in row_text:
            data_start_row = i + 1  # Данные начинаются со следующей строки
            break
        header_rows.append(row)
    
    # Удаляем строки с данными (оставляем только заголовки)
    while len(table.rows) > data_start_row:
        table._element.remove(table.rows[-1]._element)
    
    # Определяем, где находится столбец с оценкой
    # В шаблоне "Оценка" может быть в отдельной ячейке справа
    grade_col_index = None
    if table.rows and len(table.rows[0].cells) >= 5:
        grade_col_index = 4  # 5-я колонка (индекс 4)
    elif table.rows and len(table.rows[0].cells) >= 4:
        grade_col_index = 3  # 4-я колонка
    
    # Заполняем данными
    total_ze = 0
    
    for disc in transcript_disciplines:
        row = table.add_row()
        cells = row.cells
        
        # Название дисциплины
        if len(cells) >= 1:
            _cell_write(cells[0], disc.get('original_name', disc.get('name', '—')), size_pt=10)
        
        # Семестр
        if len(cells) >= 2:
            _cell_write(cells[1], _fmt_sem(disc.get('semester')), size_pt=10)
        
        # Объем ЗЕ
        if len(cells) >= 3:
            hours = disc.get('hours', 0)
            _cell_write(cells[2], _fmt_ze(hours), size_pt=10)
            try:
                total_ze += float(hours) if hours else 0
            except:
                pass
        
        # Форма контроля
        if len(cells) >= 4:
            _cell_write(cells[3], _get_control_form(disc), size_pt=10)
        
        # Оценка (последняя колонка, справа)
        if grade_col_index is not None and len(cells) > grade_col_index:
            _cell_write(cells[grade_col_index], _fmt_grade_cell(disc.get('grade'), disc.get('normalized_grade')), size_pt=10)
        elif len(cells) >= 5:
            _cell_write(cells[4], _fmt_grade_cell(disc.get('grade'), disc.get('normalized_grade')), size_pt=10)
    
    # Добавляем строку "Итого" если нужно
    # Ищем, есть ли в таблице строка "Всего"
    has_total_row = False
    for row in table.rows:
        if 'всего' in row.cells[0].text.lower() if row.cells else False:
            has_total_row = True
            break
    
    if not has_total_row:
        total_row = table.add_row()
        cells = total_row.cells
        if len(cells) >= 1:
            _cell_write(cells[0], 'Всего', size_pt=10, bold=True)
        if len(cells) >= 3:
            _cell_write(cells[2], _fmt_ze(total_ze), size_pt=10, bold=True)


def fill_table_2_comparison(table, match_results: Dict, curriculum: List[Dict]):
    """
    Таблица 2: Сопоставление (новый УП ↔ старый УП)
    Колонки: № | Наименование по новому УП | Семестр | ЗЕ | Форма контроля |
             Наименование по старому УП | Семестр | ЗЕ | Форма контроля | Итоговая оценка
    """
    if table is None:
        return
    
    # Определяем строку с заголовками (первая строка)
    # Оставляем только заголовок, удаляем все остальные строки
    while len(table.rows) > 1:
        table._element.remove(table.rows[-1]._element)
    
    row_num = 1
    total_new_ze = 0
    total_old_ze = 0
    
    # Автоматически сопоставленные
    for item in match_results.get('matched', []):
        trans_disc = item.get('transcript_discipline', {})
        curr_disc = item.get('curriculum_discipline', {})
        
        row = table.add_row()
        cells = row.cells
        
        _cell_write(cells[0], str(row_num), size_pt=9)
        _cell_write(cells[1], curr_disc.get('original_name', '—'), size_pt=9)
        _cell_write(cells[2], _fmt_sem(curr_disc.get('semester')), size_pt=9)
        
        curr_ze = curr_disc.get('hours', 0)
        _cell_write(cells[3], _fmt_ze(curr_ze), size_pt=9)
        try:
            total_new_ze += float(curr_ze) if curr_ze else 0
        except:
            pass
        
        _cell_write(cells[4], _get_control_form(curr_disc), size_pt=9)
        _cell_write(cells[5], trans_disc.get('original_name', '—'), size_pt=9)
        _cell_write(cells[6], _fmt_sem(trans_disc.get('semester')), size_pt=9)
        
        old_ze = trans_disc.get('hours', 0)
        _cell_write(cells[7], _fmt_ze(old_ze), size_pt=9)
        try:
            total_old_ze += float(old_ze) if old_ze else 0
        except:
            pass
        
        _cell_write(cells[8], _get_control_form(trans_disc), size_pt=9)
        _cell_write(cells[9], _fmt_grade_cell(trans_disc.get('grade'), trans_disc.get('normalized_grade')), size_pt=9)
        row_num += 1
    
    # Ручные сопоставления (переаттестация)
    for item in match_results.get('manual', []):
        if item.get('status') == 'manual_matched' and item.get('selected_match'):
            trans_disc = item.get('transcript_discipline', {})
            curr_disc = item.get('selected_match', {})
            
            row = table.add_row()
            cells = row.cells
            
            _cell_write(cells[0], str(row_num), size_pt=9)
            _cell_write(cells[1], curr_disc.get('original_name', '—'), size_pt=9)
            _cell_write(cells[2], _fmt_sem(curr_disc.get('semester')), size_pt=9)
            
            curr_ze = curr_disc.get('hours', 0)
            _cell_write(cells[3], _fmt_ze(curr_ze), size_pt=9)
            try:
                total_new_ze += float(curr_ze) if curr_ze else 0
            except:
                pass
            
            _cell_write(cells[4], _get_control_form(curr_disc), size_pt=9)
            _cell_write(cells[5], trans_disc.get('original_name', '—'), size_pt=9)
            _cell_write(cells[6], _fmt_sem(trans_disc.get('semester')), size_pt=9)
            
            old_ze = trans_disc.get('hours', 0)
            _cell_write(cells[7], _fmt_ze(old_ze), size_pt=9)
            try:
                total_old_ze += float(old_ze) if old_ze else 0
            except:
                pass
            
            _cell_write(cells[8], _get_control_form(trans_disc), size_pt=9)
            _cell_write(cells[9], _fmt_grade_cell(trans_disc.get('grade'), trans_disc.get('normalized_grade')), size_pt=9)
            row_num += 1
    
    # Добавляем строку "Всего"
    total_row = table.add_row()
    cells = total_row.cells
    _cell_write(cells[0], '', size_pt=9)
    _cell_write(cells[1], 'Всего', size_pt=9, bold=True)
    _cell_write(cells[3], _fmt_ze(total_new_ze), size_pt=9, bold=True)
    _cell_write(cells[7], _fmt_ze(total_old_ze), size_pt=9, bold=True)


def fill_table_3_need_study(table, need_study: List[Dict]):
    """
    Таблица 3: Перечень дисциплин или практик, подлежащих изучению
    Колонки: Название | Семестр | Объем, ЗЕ | Форма контроля | Кафедра | Объем работы преподавателя | Срок завершения
    """
    if table is None:
        return
    
    # Оставляем только заголовок (первую строку)
    while len(table.rows) > 1:
        table._element.remove(table.rows[-1]._element)
    
    total_ze = 0
    
    for disc in need_study:
        row = table.add_row()
        cells = row.cells
        
        _cell_write(cells[0], disc.get('name', '—'), size_pt=9)
        _cell_write(cells[1], _fmt_sem(disc.get('semester')), size_pt=9)
        
        hours = disc.get('hours', 0)
        _cell_write(cells[2], _fmt_ze(hours), size_pt=9)
        try:
            total_ze += float(hours) if hours else 0
        except:
            pass
        
        _cell_write(cells[3], disc.get('control_form', '—'), size_pt=9)
        _cell_write(cells[4], disc.get('department', 'ИСТ'), size_pt=9)
        
        # Объем работы преподавателя
        teacher_hours = disc.get('teacher_hours', '')
        if not teacher_hours:
            if hours:
                try:
                    h = float(hours)
                    teacher_hours = f"Л-{int(h)*4}; ПЗ-{int(h)*6}"
                except:
                    teacher_hours = '—'
            else:
                teacher_hours = '—'
        _cell_write(cells[5], teacher_hours, size_pt=9)
        
        # Срок завершения
        deadline = disc.get('deadline', '01.02.2027')
        _cell_write(cells[6], deadline, size_pt=9)
    
    # Добавляем строку "Всего"
    total_row = table.add_row()
    cells = total_row.cells
    _cell_write(cells[0], 'Всего', size_pt=9, bold=True)
    _cell_write(cells[2], _fmt_ze(total_ze), size_pt=9, bold=True)


def fill_from_template(template_path: str, 
                       transcript_disciplines: List[Dict],
                       match_results: Dict,
                       curriculum: List[Dict],
                       final_results: Dict,
                       meta: Dict) -> BytesIO:
    """
    Заполнение всех таблиц в шаблоне
    """
    doc = Document(template_path)
    
    # ===== ЗАМЕНА ТЕКСТОВЫХ ПОЛЕЙ (ФИО, группа, даты) =====
    student_full = f"{meta.get('student_genitive', '_______________')}, обучающегося по направлению подготовки {meta.get('direction', '_______________')} группы {meta.get('group', '_______________')}"
    profile_text = f"Направленность (профиль): «{meta.get('profile', '_______________')}»"
    deadline_text = f"Срок ликвидации разницы в учебных планах: {meta.get('deadline', '_______________')}"
    
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if 'обучающегося по направлению подготовки' in text:
            paragraph.clear()
            run = paragraph.add_run(student_full)
            _run_font(run, size_pt=12)
        elif 'Направленность (профиль):' in text and 'обучающегося' not in text:
            paragraph.clear()
            run = paragraph.add_run(profile_text)
            _run_font(run, size_pt=12)
        elif 'Срок ликвидации разницы' in text:
            paragraph.clear()
            run = paragraph.add_run(deadline_text)
            _run_font(run, size_pt=12)
        elif '___________' in text and 'Голкина' in text:
            paragraph.clear()
            run = paragraph.add_run(f'___________    {meta.get("prorector_fio", "В.А. Голкина")}')
            _run_font(run, size_pt=12)
    
    # ===== ЗАПОЛНЕНИЕ ТРЁХ ТАБЛИЦ =====
    # Таблица 0: Дисциплины учебного плана ЯГТУ (освоенные)
    if len(doc.tables) >= 1:
        fill_table_1_old_plan(doc.tables[0], transcript_disciplines, meta)
    
    # Таблица 1: Сопоставление (новый УП ↔ старый УП)
    if len(doc.tables) >= 2:
        fill_table_2_comparison(doc.tables[1], match_results, curriculum)
    
    # Таблица 2: Перечень дисциплин к изучению
    if len(doc.tables) >= 3:
        fill_table_3_need_study(doc.tables[2], final_results.get('need_study', []))
    
    # ===== ЗАМЕНА ДАТ И ПОДПИСЕЙ =====
    year = meta.get('sign_year', '2026')
    for paragraph in doc.paragraphs:
        if '«____» ____________ 20__ г.' in paragraph.text:
            new_text = paragraph.text.replace('20__ г.', f'{year} г.')
            paragraph.clear()
            run = paragraph.add_run(new_text)
            _run_font(run, size_pt=12)
    
    # Замена подписи студента
    student_fio = meta.get('student_sign_fio', '_______________')
    for paragraph in doc.paragraphs:
        if '_________________' in paragraph.text and 'Егоров' in paragraph.text:
            new_text = paragraph.text.replace('_________________', student_fio)
            # Заменяем также в строке с обучающимся
            if 'Обучающийся' in new_text:
                lines = new_text.split('\n')
                for i, line in enumerate(lines):
                    if '_________________' in line:
                        lines[i] = line.replace('_________________', student_fio)
                new_text = '\n'.join(lines)
            paragraph.clear()
            run = paragraph.add_run(new_text)
            _run_font(run, size_pt=12)
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ============================================================================
# ПРОГРАММНАЯ ГЕНЕРАЦИЯ (ЕСЛИ ШАБЛОНА НЕТ)
# ============================================================================

def _build_programmatic(transcript_disciplines: List[Dict],
                        match_results: Dict,
                        curriculum: List[Dict],
                        final_results: Dict,
                        meta: Dict) -> BytesIO:
    """Полная программная генерация документа"""
    m = {**DEFAULT_PLAN_META, **(meta or {})}
    doc = Document()
    
    # Настройка полей
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    
    # Шапка с утверждением
    approve_lines = [
        'УТВЕРЖДАЮ',
        'Проректор по',
        'образовательной',
        'деятельности и',
        'воспитательной работе',
        f'___________    {m.get("prorector_fio", "В.А. Голкина")}',
        f'«____» ____________ {m.get("sign_year", "2026")} г.',
    ]
    for line in approve_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run_font(p.add_run(line), size_pt=12)
    
    doc.add_paragraph()
    
    # Заголовок
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_font(p.add_run('Индивидуальный план\nликвидации разницы в учебных планах'), size_pt=14, bold=True)
    
    doc.add_paragraph()
    
    # Информация о студенте
    _para(doc, f'{m["student_genitive"]}, обучающегося по направлению подготовки {m["direction"]} группы {m["group"]}', size_pt=12)
    _para(doc, f'Направленность (профиль): «{m["profile"]}»', size_pt=12)
    _para(doc, f'Срок ликвидации разницы в учебных планах: {m["deadline"]}', size_pt=12)
    
    doc.add_paragraph()
    
    # ===== ТАБЛИЦА 1 =====
    # Заголовок таблицы 1
    p1 = doc.add_paragraph()
    _run_font(p1.add_run('Дисциплины учебного плана ЯГТУ'), size_pt=11, bold=True)
    p2 = doc.add_paragraph()
    _run_font(p2.add_run(f'направления подготовки {m["direction"]}'), size_pt=10)
    p3 = doc.add_paragraph()
    _run_font(p3.add_run('Институт цифровых систем'), size_pt=10)
    
    # Пустая строка для отступа
    doc.add_paragraph()
    
    headers1 = ['Название дисциплины/практики', 'Семестр', 'Объем, ЗЕ', 'Форма контроля', 'Оценка']
    table1 = doc.add_table(rows=1 + len(transcript_disciplines) + 1, cols=5)  # +1 для строки "Всего"
    table1.style = 'Table Grid'
    
    for j, h in enumerate(headers1):
        _cell_write(table1.rows[0].cells[j], h, size_pt=10, bold=True)
    
    total_ze = 0
    for i, disc in enumerate(transcript_disciplines, 1):
        row = table1.rows[i].cells
        _cell_write(row[0], disc.get('original_name', disc.get('name', '—')), size_pt=9)
        _cell_write(row[1], _fmt_sem(disc.get('semester')), size_pt=9)
        hours = disc.get('hours', 0)
        _cell_write(row[2], _fmt_ze(hours), size_pt=9)
        try:
            total_ze += float(hours) if hours else 0
        except:
            pass
        _cell_write(row[3], _get_control_form(disc), size_pt=9)
        _cell_write(row[4], _fmt_grade_cell(disc.get('grade'), disc.get('normalized_grade')), size_pt=9)
    
    # Строка "Всего"
    total_row = table1.rows[len(transcript_disciplines) + 1].cells
    _cell_write(total_row[0], 'Всего', size_pt=9, bold=True)
    _cell_write(total_row[2], _fmt_ze(total_ze), size_pt=9, bold=True)
    
    doc.add_paragraph()
    
    # ===== ТАБЛИЦА 2 =====
    headers2 = ['№', 'Наименование дисциплин по новому уч. плану', 'Семестр', 'Объем, ЗЕ', 'Форма контроля',
                'Наименование дисциплин по старому уч. плану', 'Семестр', 'Объем, ЗЕ', 'Форма контроля', 'Итоговая оценка']
    
    matched_items = []
    for item in match_results.get('matched', []):
        matched_items.append({
            'trans': item.get('transcript_discipline', {}),
            'curr': item.get('curriculum_discipline', {})
        })
    for item in match_results.get('manual', []):
        if item.get('status') == 'manual_matched' and item.get('selected_match'):
            matched_items.append({
                'trans': item.get('transcript_discipline', {}),
                'curr': item.get('selected_match', {})
            })
    
    table2 = doc.add_table(rows=1 + len(matched_items) + 1, cols=10)  # +1 для строки "Всего"
    table2.style = 'Table Grid'
    
    for j, h in enumerate(headers2):
        _cell_write(table2.rows[0].cells[j], h, size_pt=9, bold=True)
    
    total_new_ze = 0
    total_old_ze = 0
    
    for i, item in enumerate(matched_items, 1):
        row = table2.rows[i].cells
        trans = item['trans']
        curr = item['curr']
        
        _cell_write(row[0], str(i), size_pt=9)
        _cell_write(row[1], curr.get('original_name', '—'), size_pt=9)
        _cell_write(row[2], _fmt_sem(curr.get('semester')), size_pt=9)
        
        curr_ze = curr.get('hours', 0)
        _cell_write(row[3], _fmt_ze(curr_ze), size_pt=9)
        try:
            total_new_ze += float(curr_ze) if curr_ze else 0
        except:
            pass
        
        _cell_write(row[4], _get_control_form(curr), size_pt=9)
        _cell_write(row[5], trans.get('original_name', '—'), size_pt=9)
        _cell_write(row[6], _fmt_sem(trans.get('semester')), size_pt=9)
        
        old_ze = trans.get('hours', 0)
        _cell_write(row[7], _fmt_ze(old_ze), size_pt=9)
        try:
            total_old_ze += float(old_ze) if old_ze else 0
        except:
            pass
        
        _cell_write(row[8], _get_control_form(trans), size_pt=9)
        _cell_write(row[9], _fmt_grade_cell(trans.get('grade'), trans.get('normalized_grade')), size_pt=9)
    
    # Строка "Всего"
    total_row2 = table2.rows[len(matched_items) + 1].cells
    _cell_write(total_row2[1], 'Всего', size_pt=9, bold=True)
    _cell_write(total_row2[3], _fmt_ze(total_new_ze), size_pt=9, bold=True)
    _cell_write(total_row2[7], _fmt_ze(total_old_ze), size_pt=9, bold=True)
    
    doc.add_paragraph()
    
    # ===== ТАБЛИЦА 3 =====
    _para(doc, 'Перечень дисциплин или практик подлежащих изучению или прохождению', bold=True, size_pt=11)
    
    headers3 = ['Название дисциплины/практики', 'Семестр', 'Объем, ЗЕ', 'Форма контроля', 'Кафедра', 'Объем работы преподавателя, а.ч.', 'Срок завершения']
    need_study_list = final_results.get('need_study', [])
    table3 = doc.add_table(rows=1 + len(need_study_list) + 1, cols=7)  # +1 для строки "Всего"
    table3.style = 'Table Grid'
    
    for j, h in enumerate(headers3):
        _cell_write(table3.rows[0].cells[j], h, size_pt=9, bold=True)
    
    total_study_ze = 0
    for i, disc in enumerate(need_study_list, 1):
        row = table3.rows[i].cells
        _cell_write(row[0], disc.get('name', '—'), size_pt=9)
        _cell_write(row[1], _fmt_sem(disc.get('semester')), size_pt=9)
        
        hours = disc.get('hours', 0)
        _cell_write(row[2], _fmt_ze(hours), size_pt=9)
        try:
            total_study_ze += float(hours) if hours else 0
        except:
            pass
        
        _cell_write(row[3], disc.get('control_form', '—'), size_pt=9)
        _cell_write(row[4], disc.get('department', 'ИСТ'), size_pt=9)
        
        teacher_hours = disc.get('teacher_hours', '')
        if not teacher_hours:
            if hours:
                try:
                    h = float(hours)
                    teacher_hours = f"Л-{int(h)*4}; ПЗ-{int(h)*6}"
                except:
                    teacher_hours = '—'
            else:
                teacher_hours = '—'
        _cell_write(row[5], teacher_hours, size_pt=9)
        _cell_write(row[6], disc.get('deadline', '01.02.2027'), size_pt=9)
    
    # Строка "Всего"
    total_row3 = table3.rows[len(need_study_list) + 1].cells
    _cell_write(total_row3[0], 'Всего', size_pt=9, bold=True)
    _cell_write(total_row3[2], _fmt_ze(total_study_ze), size_pt=9, bold=True)
    
    doc.add_paragraph()
    
    # ===== ПОДПИСИ =====
    doc.add_paragraph()
    _para(doc, 'СОГЛАСОВАНО:', bold=True, size_pt=12)
    doc.add_paragraph()
    
    sig_block = (f'Обучающийся\n'
                 f'«____» ____________ {m["sign_year"]} г.\t\t_________________\t\t{m.get("student_sign_fio", "_______________")}\n\n'
                 f'{m.get("deputy_title", "Заместитель директора Института цифровых систем")}\n'
                 f'«____» ____________ {m["sign_year"]} г.\t\t_________________\t\t{m.get("deputy_fio", "_______________")}')
    _para(doc, sig_block, size_pt=12)
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ============================================================================
# ОСНОВНАЯ ФУНКЦИЯ
# ============================================================================

def build_individual_plan_docx(final_results: Dict[str, Any],
                                transcript_disciplines: Optional[List[Dict]] = None,
                                match_results: Optional[Dict[str, Any]] = None,
                                curriculum: Optional[List[Dict]] = None,
                                meta: Optional[Dict[str, Any]] = None) -> BytesIO:
    """
    Создание документа индивидуального плана
    """
    meta = meta or {}
    transcript_disciplines = transcript_disciplines or []
    match_results = match_results or {}
    curriculum = curriculum or []
    
    template_path = 'plan_template.docx'
    if os.path.exists(template_path):
        try:
            return fill_from_template(template_path, transcript_disciplines, match_results, curriculum, final_results, meta)
        except Exception as e:
            print(f"Ошибка шаблона: {e}")
            import traceback
            traceback.print_exc()
    
    return _build_programmatic(transcript_disciplines, match_results, curriculum, final_results, meta)